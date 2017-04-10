from docx import Document
import os


PRIZE = {'阿才阿聪睇广州': '一', '舌尖上的广州': '二', '二狗子成都寻爱记': '二', '一万零四百八十一公里': '二'}

namelist = os.listdir('./namelist/')
doc = Document('mb.docx')
count = 0
for filename in namelist:
    title = filename.split('.')[0]
    docc = Document('./namelist/' + filename)
    teamnamelist = list()
    for cell in docc.tables[0].rows[7].cells:
        text = cell.text.replace(' ', '')
        if text == '姓名' or text == '':
            continue
        teamnamelist.append(cell.text)
    prize = PRIZE.get(title, '三')
    for name in teamnamelist:
        count += 1
        doc.save(str(count) + '.docx')
        doccc = Document(str(count) + '.docx')
        doccc.paragraphs[0].runs[0].text = doccc.paragraphs[
            0].runs[0].text.replace('占位符', name)
        doccc.paragraphs[1].runs[2].text = doccc.paragraphs[
            1].runs[2].text.replace('占位符', '《' + title + '》')
        doccc.paragraphs[1].runs[4].text = doccc.paragraphs[
            1].runs[4].text.replace('占符', prize)
        doccc.save(str(count) + '.docx')
